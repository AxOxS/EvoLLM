"""RAG document management endpoints – upload, list, delete."""

from __future__ import annotations

import io

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from sqlalchemy.orm import Session

from backend.database import get_db
from backend.models import User, Document
from backend.schemas import DocumentResponse
from backend.auth import get_current_user
from backend.services.rag_service import rag_service

router = APIRouter(prefix="/rag", tags=["rag"])


def _extract_text(filename: str, content: bytes) -> str:
    """Extract plain text from uploaded file."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"

    if ext == "pdf":
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)

    if ext == "docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs if p.text)

    # Default: treat as plain text
    return content.decode("utf-8", errors="replace")


@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    content = await file.read()
    if not content:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Empty file")

    filename = file.filename or "unknown.txt"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    if ext not in ("pdf", "txt", "docx"):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported file type. Use PDF, TXT, or DOCX.")

    text = _extract_text(filename, content)
    size_kb = round(len(content) / 1024, 1)

    doc = Document(
        user_id=user.id,
        filename=filename,
        content_text=text,
        doc_type=ext,
        size_kb=size_kb,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    # Add to ChromaDB vector store
    rag_service.add_document(doc.id, text, {"filename": filename, "user_id": user.id})
    doc.chromadb_id = doc.id
    db.commit()

    return doc


@router.get("/documents", response_model=list[DocumentResponse])
def list_documents(user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    docs = (
        db.query(Document)
        .filter(Document.user_id == user.id)
        .order_by(Document.uploaded_at.desc())
        .all()
    )
    return docs


@router.delete("/documents/{doc_id}")
def delete_document(doc_id: str, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == user.id).first()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    # Remove from ChromaDB
    rag_service.delete_document(doc.id)

    db.delete(doc)
    db.commit()
    return {"ok": True}
